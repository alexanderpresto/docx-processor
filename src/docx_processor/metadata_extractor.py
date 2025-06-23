"""
Enhanced metadata extraction module for docx-processor v2.0.

This module provides comprehensive metadata extraction from DOCX files including:
- Document properties (core, extended, custom)
- Author and contributor information
- Timestamps and document statistics
- Comments and revision tracking
- Language and locale information

Uses both python-docx for structured access and direct XML parsing for
comprehensive metadata extraction.
"""

import os
import json
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class MetadataExtractor:
    """Extracts comprehensive metadata from DOCX files."""
    
    def __init__(self):
        """Initialize the metadata extractor."""
        # XML namespaces used in DOCX files
        self.namespaces = {
            'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
            'dc': 'http://purl.org/dc/elements/1.1/',
            'dcterms': 'http://purl.org/dc/terms/',
            'dcmitype': 'http://purl.org/dc/dcmitype/',
            'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
            'app': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties',
            'vt': 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
    
    def extract_all_metadata(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract all available metadata from a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary containing all extracted metadata
        """
        metadata = {
            'file_info': self._get_file_info(docx_path),
            'core_properties': {},
            'extended_properties': {},
            'custom_properties': {},
            'comments': [],
            'revision_info': {},
            'document_statistics': {},
            'language_info': {},
            'extraction_timestamp': datetime.now().isoformat()
        }
        
        try:
            # Try python-docx first for structured access
            doc = Document(docx_path)
            metadata.update(self._extract_with_python_docx(doc))
            
            # Enhance with direct XML parsing
            metadata.update(self._extract_with_xml_parsing(docx_path))
            
        except (PackageNotFoundError, Exception) as e:
            print(f"Warning: Could not extract metadata with python-docx: {e}")
            # Fallback to XML-only extraction
            try:
                metadata.update(self._extract_with_xml_parsing(docx_path))
            except Exception as xml_error:
                print(f"Warning: XML metadata extraction also failed: {xml_error}")
                metadata['extraction_errors'] = [str(e), str(xml_error)]
        
        return metadata
    
    def _get_file_info(self, docx_path: str) -> Dict[str, Any]:
        """Extract basic file information."""
        try:
            stat = os.stat(docx_path)
            return {
                'filename': os.path.basename(docx_path),
                'filepath': os.path.abspath(docx_path),
                'file_size_bytes': stat.st_size,
                'file_size_mb': round(stat.st_size / (1024 * 1024), 2),
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed': datetime.fromtimestamp(stat.st_atime).isoformat()
            }
        except Exception as e:
            return {'error': f"Could not extract file info: {e}"}
    
    def _extract_with_python_docx(self, doc: Document) -> Dict[str, Any]:
        """Extract metadata using python-docx library."""
        metadata = {}
        
        try:
            # Core properties
            core_props = doc.core_properties
            metadata['core_properties'] = {
                'title': core_props.title,
                'subject': core_props.subject,
                'creator': core_props.author,
                'keywords': core_props.keywords,
                'description': core_props.comments,
                'category': core_props.category,
                'language': core_props.language,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
                'version': core_props.version
            }
            
            # Document statistics
            metadata['document_statistics'] = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections)
            }
            
            # Extract section information
            sections_info = []
            for i, section in enumerate(doc.sections):
                section_info = {
                    'section_number': i + 1,
                    'start_type': str(section.start_type),
                    'orientation': str(section.orientation),
                    'page_width': section.page_width.inches if section.page_width else None,
                    'page_height': section.page_height.inches if section.page_height else None,
                    'left_margin': section.left_margin.inches if section.left_margin else None,
                    'right_margin': section.right_margin.inches if section.right_margin else None,
                    'top_margin': section.top_margin.inches if section.top_margin else None,
                    'bottom_margin': section.bottom_margin.inches if section.bottom_margin else None
                }
                sections_info.append(section_info)
            
            metadata['sections'] = sections_info
            
        except Exception as e:
            metadata['python_docx_error'] = str(e)
        
        return metadata
    
    def _extract_with_xml_parsing(self, docx_path: str) -> Dict[str, Any]:
        """Extract metadata through direct XML parsing of DOCX internals."""
        metadata = {}
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Extract core properties
                core_props = self._parse_core_properties(zip_ref)
                if core_props:
                    metadata['core_properties'].update(core_props)
                
                # Extract extended properties
                extended_props = self._parse_extended_properties(zip_ref)
                if extended_props:
                    metadata['extended_properties'] = extended_props
                
                # Extract custom properties
                custom_props = self._parse_custom_properties(zip_ref)
                if custom_props:
                    metadata['custom_properties'] = custom_props
                
                # Extract comments
                comments = self._parse_comments(zip_ref)
                if comments:
                    metadata['comments'] = comments
                
                # Extract revision information
                revision_info = self._parse_revision_info(zip_ref)
                if revision_info:
                    metadata['revision_info'] = revision_info
                
        except Exception as e:
            metadata['xml_parsing_error'] = str(e)
        
        return metadata
    
    def _parse_core_properties(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse core properties from docProps/core.xml."""
        try:
            if 'docProps/core.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('docProps/core.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                properties = {}
                
                # Map XML elements to property names
                element_map = {
                    'title': 'dc:title',
                    'subject': 'dc:subject',
                    'creator': 'dc:creator',
                    'keywords': 'cp:keywords',
                    'description': 'dc:description',
                    'category': 'cp:category',
                    'language': 'dc:language',
                    'created': 'dcterms:created',
                    'modified': 'dcterms:modified',
                    'last_modified_by': 'cp:lastModifiedBy',
                    'revision': 'cp:revision',
                    'version': 'cp:version'
                }
                
                for prop_name, xml_path in element_map.items():
                    element = root.find(xml_path, self.namespaces)
                    if element is not None:
                        properties[prop_name] = element.text
                
                return properties
                
        except Exception as e:
            print(f"Warning: Could not parse core properties: {e}")
            return None
    
    def _parse_extended_properties(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse extended properties from docProps/app.xml."""
        try:
            if 'docProps/app.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('docProps/app.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                properties = {}
                
                # Map XML elements to property names
                element_map = {
                    'application': 'app:Application',
                    'app_version': 'app:AppVersion',
                    'document_security': 'app:DocSecurity',
                    'scale_crop': 'app:ScaleCrop',
                    'company': 'app:Company',
                    'links_up_to_date': 'app:LinksUpToDate',
                    'shared_doc': 'app:SharedDoc',
                    'hyperlinks_changed': 'app:HyperlinksChanged',
                    'template': 'app:Template',
                    'total_time': 'app:TotalTime',
                    'pages': 'app:Pages',
                    'words': 'app:Words',
                    'characters': 'app:Characters',
                    'characters_with_spaces': 'app:CharactersWithSpaces',
                    'lines': 'app:Lines',
                    'paragraphs': 'app:Paragraphs'
                }
                
                for prop_name, xml_path in element_map.items():
                    element = root.find(xml_path, self.namespaces)
                    if element is not None:
                        # Convert numeric values
                        if prop_name in ['pages', 'words', 'characters', 'characters_with_spaces', 
                                       'lines', 'paragraphs', 'total_time']:
                            try:
                                properties[prop_name] = int(element.text) if element.text else 0
                            except ValueError:
                                properties[prop_name] = element.text
                        else:
                            properties[prop_name] = element.text
                
                return properties
                
        except Exception as e:
            print(f"Warning: Could not parse extended properties: {e}")
            return None
    
    def _parse_custom_properties(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse custom properties from docProps/custom.xml."""
        try:
            if 'docProps/custom.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('docProps/custom.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                properties = {}
                
                # Custom properties are stored differently
                for prop in root.findall('.//property', self.namespaces):
                    name = prop.get('name')
                    if name:
                        # Find the value element
                        for child in prop:
                            if child.text:
                                properties[name] = child.text
                                break
                
                return properties
                
        except Exception as e:
            print(f"Warning: Could not parse custom properties: {e}")
            return None
    
    def _parse_comments(self, zip_ref: zipfile.ZipFile) -> Optional[List[Dict[str, Any]]]:
        """Parse comments from word/comments.xml."""
        try:
            if 'word/comments.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('word/comments.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                comments = []
                
                for comment in root.findall('.//w:comment', self.namespaces):
                    comment_data = {
                        'id': comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id'),
                        'author': comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author'),
                        'date': comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date'),
                        'text': ''
                    }
                    
                    # Extract comment text
                    text_parts = []
                    for text_elem in comment.findall('.//w:t', self.namespaces):
                        if text_elem.text:
                            text_parts.append(text_elem.text)
                    
                    comment_data['text'] = ''.join(text_parts)
                    comments.append(comment_data)
                
                return comments
                
        except Exception as e:
            print(f"Warning: Could not parse comments: {e}")
            return None
    
    def _parse_revision_info(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse revision tracking information from word/document.xml."""
        try:
            if 'word/document.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('word/document.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                revision_info = {
                    'has_revisions': False,
                    'revision_count': 0,
                    'tracked_changes': []
                }
                
                # Look for revision tracking elements
                revisions = root.findall('.//w:ins', self.namespaces) + \
                           root.findall('.//w:del', self.namespaces)
                
                if revisions:
                    revision_info['has_revisions'] = True
                    revision_info['revision_count'] = len(revisions)
                    
                    for rev in revisions[:10]:  # Limit to first 10 revisions
                        rev_data = {
                            'type': rev.tag.split('}')[-1] if '}' in rev.tag else rev.tag,
                            'author': rev.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author'),
                            'date': rev.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date'),
                            'id': rev.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        }
                        revision_info['tracked_changes'].append(rev_data)
                
                return revision_info
                
        except Exception as e:
            print(f"Warning: Could not parse revision info: {e}")
            return None
    
    def save_metadata(self, metadata: Dict[str, Any], output_path: str) -> None:
        """
        Save extracted metadata to a JSON file.
        
        Args:
            metadata: The metadata dictionary to save
            output_path: Path where to save the metadata JSON file
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False, default=str)
            print(f"Metadata saved to {output_path}")
        except Exception as e:
            print(f"Error saving metadata: {e}")
    
    def get_metadata_summary(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Create a summary of the most important metadata fields.
        
        Args:
            metadata: Full metadata dictionary
            
        Returns:
            Summary dictionary with key metadata fields
        """
        summary = {}
        
        # File information
        if 'file_info' in metadata:
            summary['file'] = {
                'name': metadata['file_info'].get('filename'),
                'size_mb': metadata['file_info'].get('file_size_mb'),
                'modified': metadata['file_info'].get('modified')
            }
        
        # Core properties
        if 'core_properties' in metadata:
            core = metadata['core_properties']
            summary['document'] = {
                'title': core.get('title'),
                'author': core.get('creator'),
                'subject': core.get('subject'),
                'created': core.get('created'),
                'modified': core.get('modified'),
                'language': core.get('language')
            }
        
        # Statistics
        if 'extended_properties' in metadata:
            ext = metadata['extended_properties']
            summary['statistics'] = {
                'pages': ext.get('pages'),
                'words': ext.get('words'),
                'characters': ext.get('characters'),
                'paragraphs': ext.get('paragraphs')
            }
        
        # Comments and revisions
        summary['content_features'] = {
            'has_comments': len(metadata.get('comments', [])) > 0,
            'comment_count': len(metadata.get('comments', [])),
            'has_revisions': metadata.get('revision_info', {}).get('has_revisions', False),
            'custom_properties': len(metadata.get('custom_properties', {})) > 0
        }
        
        return summary