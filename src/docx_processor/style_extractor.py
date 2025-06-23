"""
Style information extraction module for docx-processor v2.0.

This module extracts comprehensive style and formatting information from DOCX files including:
- Font families and sizes used throughout the document
- Paragraph and character styles mapping
- Color palette extraction
- Layout information (margins, orientation, page size)
- Header and footer content
- Section breaks and formatting
- Theme information

Uses both python-docx for structured access and direct XML parsing for
complete style information extraction.
"""

import os
import json
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, List, Optional, Set, Tuple
from collections import defaultdict, Counter
from docx import Document
from docx.shared import RGBColor
from docx.opc.exceptions import PackageNotFoundError


class StyleExtractor:
    """Extracts comprehensive style and formatting information from DOCX files."""
    
    def __init__(self):
        """Initialize the style extractor."""
        # XML namespaces used in DOCX files
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'v': 'urn:schemas-microsoft-com:vml'
        }
    
    def extract_all_styles(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract all style information from a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary containing all extracted style information
        """
        styles = {
            'fonts': {},
            'colors': {},
            'paragraph_styles': {},
            'character_styles': {},
            'table_styles': {},
            'numbering_styles': {},
            'layout_info': {},
            'headers_footers': {},
            'theme_info': {},
            'style_usage_stats': {},
            'extraction_timestamp': None
        }
        
        try:
            # Try python-docx first for structured access
            doc = Document(docx_path)
            styles.update(self._extract_with_python_docx(doc))
            
            # Enhance with direct XML parsing
            styles.update(self._extract_with_xml_parsing(docx_path))
            
        except (PackageNotFoundError, Exception) as e:
            print(f"Warning: Could not extract styles with python-docx: {e}")
            # Fallback to XML-only extraction
            try:
                styles.update(self._extract_with_xml_parsing(docx_path))
            except Exception as xml_error:
                print(f"Warning: XML style extraction also failed: {xml_error}")
                styles['extraction_errors'] = [str(e), str(xml_error)]
        
        return styles
    
    def _extract_with_python_docx(self, doc: Document) -> Dict[str, Any]:
        """Extract style information using python-docx library."""
        styles = {}
        
        try:
            # Document styles
            if hasattr(doc, 'styles'):
                styles['paragraph_styles'] = self._extract_paragraph_styles(doc)
                styles['character_styles'] = self._extract_character_styles(doc)
                styles['table_styles'] = self._extract_table_styles(doc)
            
            # Layout information from sections
            styles['layout_info'] = self._extract_layout_info(doc)
            
            # Font and color usage analysis
            styles['fonts'], styles['colors'] = self._analyze_text_formatting(doc)
            
            # Headers and footers
            styles['headers_footers'] = self._extract_headers_footers(doc)
            
        except Exception as e:
            styles['python_docx_style_error'] = str(e)
        
        return styles
    
    def _extract_paragraph_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract paragraph style definitions."""
        paragraph_styles = {}
        
        try:
            from docx.enum.style import WD_STYLE_TYPE
            
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_info = {
                        'name': style.name,
                        'style_id': getattr(style, 'style_id', None),
                        'builtin': style.builtin,
                        'hidden': style.hidden,
                        'quick_style': style.quick_style,
                        'priority': style.priority,
                        'locked': style.locked
                    }
                    
                    # Extract paragraph formatting
                    if hasattr(style, 'paragraph_format'):
                        pf = style.paragraph_format
                        style_info['paragraph_format'] = {
                            'alignment': str(pf.alignment) if pf.alignment else None,
                            'first_line_indent': self._convert_length(pf.first_line_indent),
                            'left_indent': self._convert_length(pf.left_indent),
                            'right_indent': self._convert_length(pf.right_indent),
                            'space_before': self._convert_length(pf.space_before),
                            'space_after': self._convert_length(pf.space_after),
                            'line_spacing': getattr(pf, 'line_spacing', None),
                            'keep_together': getattr(pf, 'keep_together', None),
                            'keep_with_next': getattr(pf, 'keep_with_next', None),
                            'page_break_before': getattr(pf, 'page_break_before', None),
                            'widow_control': getattr(pf, 'widow_control', None)
                        }
                    
                    # Extract font formatting
                    if hasattr(style, 'font'):
                        style_info['font'] = self._extract_font_info(style.font)
                    
                    paragraph_styles[style.name] = style_info
                    
        except Exception as e:
            print(f"Warning: Could not extract paragraph styles: {e}")
        
        return paragraph_styles
    
    def _extract_character_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract character style definitions."""
        character_styles = {}
        
        try:
            from docx.enum.style import WD_STYLE_TYPE
            
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.CHARACTER:
                    style_info = {
                        'name': style.name,
                        'style_id': getattr(style, 'style_id', None),
                        'builtin': style.builtin,
                        'hidden': style.hidden,
                        'quick_style': style.quick_style,
                        'priority': style.priority,
                        'locked': style.locked
                    }
                    
                    # Extract font formatting
                    if hasattr(style, 'font'):
                        style_info['font'] = self._extract_font_info(style.font)
                    
                    character_styles[style.name] = style_info
                    
        except Exception as e:
            print(f"Warning: Could not extract character styles: {e}")
        
        return character_styles
    
    def _extract_table_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract table style definitions."""
        table_styles = {}
        
        try:
            from docx.enum.style import WD_STYLE_TYPE
            
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.TABLE:
                    style_info = {
                        'name': style.name,
                        'style_id': getattr(style, 'style_id', None),
                        'builtin': style.builtin,
                        'hidden': style.hidden,
                        'quick_style': style.quick_style,
                        'priority': style.priority,
                        'locked': style.locked
                    }
                    
                    table_styles[style.name] = style_info
                    
        except Exception as e:
            print(f"Warning: Could not extract table styles: {e}")
        
        return table_styles
    
    def _extract_font_info(self, font) -> Dict[str, Any]:
        """Extract font formatting information."""
        font_info = {}
        
        try:
            font_info = {
                'name': font.name,
                'size': font.size.pt if font.size else None,
                'bold': font.bold,
                'italic': font.italic,
                'underline': str(font.underline) if font.underline else None,
                'strike': font.strike,
                'double_strike': font.double_strike,
                'all_caps': font.all_caps,
                'small_caps': font.small_caps,
                'shadow': font.shadow,
                'outline': font.outline,
                'rtl': font.rtl,
                'cs_bold': font.cs_bold,
                'cs_italic': font.cs_italic,
                'color': self._extract_color_info(font.color) if font.color else None,
                'highlight_color': str(font.highlight_color) if font.highlight_color else None
            }
        except Exception as e:
            font_info['extraction_error'] = str(e)
        
        return font_info
    
    def _extract_color_info(self, color) -> Optional[Dict[str, Any]]:
        """Extract color information."""
        try:
            color_info = {}
            
            if hasattr(color, 'rgb') and color.rgb:
                rgb = color.rgb
                color_info['rgb'] = f"#{rgb:06x}"
                color_info['rgb_values'] = [
                    (rgb >> 16) & 0xFF,  # Red
                    (rgb >> 8) & 0xFF,   # Green
                    rgb & 0xFF           # Blue
                ]
            
            if hasattr(color, 'theme_color') and color.theme_color:
                color_info['theme_color'] = str(color.theme_color)
            
            return color_info if color_info else None
            
        except Exception:
            return None
    
    def _convert_length(self, length) -> Optional[float]:
        """Convert docx length to points."""
        try:
            return length.pt if length else None
        except Exception:
            return None
    
    def _extract_layout_info(self, doc: Document) -> Dict[str, Any]:
        """Extract layout and page setup information."""
        layout_info = {}
        
        try:
            sections_info = []
            for i, section in enumerate(doc.sections):
                section_info = {
                    'section_number': i + 1,
                    'start_type': str(section.start_type),
                    'orientation': str(section.orientation),
                    'page_width_inches': section.page_width.inches if section.page_width else None,
                    'page_height_inches': section.page_height.inches if section.page_height else None,
                    'left_margin_inches': section.left_margin.inches if section.left_margin else None,
                    'right_margin_inches': section.right_margin.inches if section.right_margin else None,
                    'top_margin_inches': section.top_margin.inches if section.top_margin else None,
                    'bottom_margin_inches': section.bottom_margin.inches if section.bottom_margin else None,
                    'header_distance_inches': section.header_distance.inches if section.header_distance else None,
                    'footer_distance_inches': section.footer_distance.inches if section.footer_distance else None,
                    'gutter_inches': section.gutter.inches if section.gutter else None,
                    'different_first_page_header_footer': section.different_first_page_header_footer
                }
                sections_info.append(section_info)
            
            layout_info['sections'] = sections_info
            
        except Exception as e:
            layout_info['extraction_error'] = str(e)
        
        return layout_info
    
    def _analyze_text_formatting(self, doc: Document) -> Tuple[Dict[str, Any], Dict[str, Any]]:
        """Analyze font and color usage throughout the document."""
        font_usage = defaultdict(int)
        color_usage = defaultdict(int)
        font_sizes = defaultdict(int)
        
        try:
            # Analyze paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name:
                        font_usage[run.font.name] += 1
                    
                    if run.font.size:
                        font_sizes[run.font.size.pt] += 1
                    
                    if run.font.color and run.font.color.rgb:
                        color_hex = f"#{run.font.color.rgb:06x}"
                        color_usage[color_hex] += 1
            
            # Analyze tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.font.name:
                                    font_usage[run.font.name] += 1
                                
                                if run.font.size:
                                    font_sizes[run.font.size.pt] += 1
                                
                                if run.font.color and run.font.color.rgb:
                                    color_hex = f"#{run.font.color.rgb:06x}"
                                    color_usage[color_hex] += 1
            
        except Exception as e:
            print(f"Warning: Error analyzing text formatting: {e}")
        
        # Compile font information
        fonts = {
            'usage_count': dict(font_usage),
            'most_used_fonts': [font for font, count in Counter(font_usage).most_common(10)],
            'font_sizes_used': dict(font_sizes),
            'most_common_sizes': [size for size, count in Counter(font_sizes).most_common(10)]
        }
        
        # Compile color information
        colors = {
            'usage_count': dict(color_usage),
            'color_palette': list(color_usage.keys()),
            'most_used_colors': [color for color, count in Counter(color_usage).most_common(10)]
        }
        
        return fonts, colors
    
    def _extract_headers_footers(self, doc: Document) -> Dict[str, Any]:
        """Extract header and footer information."""
        headers_footers = {
            'headers': [],
            'footers': []
        }
        
        try:
            for i, section in enumerate(doc.sections):
                section_headers = {
                    'section_number': i + 1,
                    'header_text': section.header.paragraphs[0].text if section.header.paragraphs else '',
                    'first_page_header_text': (section.first_page_header.paragraphs[0].text 
                                             if section.first_page_header.paragraphs else ''),
                    'even_page_header_text': (section.even_page_header.paragraphs[0].text 
                                            if section.even_page_header.paragraphs else '')
                }
                
                section_footers = {
                    'section_number': i + 1,
                    'footer_text': section.footer.paragraphs[0].text if section.footer.paragraphs else '',
                    'first_page_footer_text': (section.first_page_footer.paragraphs[0].text 
                                             if section.first_page_footer.paragraphs else ''),
                    'even_page_footer_text': (section.even_page_footer.paragraphs[0].text 
                                            if section.even_page_footer.paragraphs else '')
                }
                
                headers_footers['headers'].append(section_headers)
                headers_footers['footers'].append(section_footers)
                
        except Exception as e:
            headers_footers['extraction_error'] = str(e)
        
        return headers_footers
    
    def _extract_with_xml_parsing(self, docx_path: str) -> Dict[str, Any]:
        """Extract style information through direct XML parsing."""
        styles = {}
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Parse styles.xml
                styles_xml = self._parse_styles_xml(zip_ref)
                if styles_xml:
                    styles.update(styles_xml)
                
                # Parse theme information
                theme_info = self._parse_theme_xml(zip_ref)
                if theme_info:
                    styles['theme_info'] = theme_info
                
                # Parse numbering styles
                numbering_styles = self._parse_numbering_xml(zip_ref)
                if numbering_styles:
                    styles['numbering_styles'] = numbering_styles
                
        except Exception as e:
            styles['xml_parsing_error'] = str(e)
        
        return styles
    
    def _parse_styles_xml(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse word/styles.xml for detailed style definitions."""
        try:
            if 'word/styles.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('word/styles.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                styles_data = {
                    'default_styles': {},
                    'custom_styles': {},
                    'style_count': 0
                }
                
                # Count total styles
                all_styles = root.findall('.//w:style', self.namespaces)
                styles_data['style_count'] = len(all_styles)
                
                # Parse individual styles
                for style in all_styles:
                    style_id = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                    style_type = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    
                    if style_id:
                        style_info = {
                            'style_id': style_id,
                            'type': style_type,
                            'default': style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}default') == '1'
                        }
                        
                        # Get style name
                        name_elem = style.find('.//w:name', self.namespaces)
                        if name_elem is not None:
                            style_info['name'] = name_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        
                        # Store in appropriate category
                        if style_info['default']:
                            styles_data['default_styles'][style_id] = style_info
                        else:
                            styles_data['custom_styles'][style_id] = style_info
                
                return styles_data
                
        except Exception as e:
            print(f"Warning: Could not parse styles.xml: {e}")
            return None
    
    def _parse_theme_xml(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse word/theme/theme1.xml for theme information."""
        try:
            theme_files = [name for name in zip_ref.namelist() if name.startswith('word/theme/')]
            if not theme_files:
                return None
            
            with zip_ref.open(theme_files[0]) as xml_file:
                root = ET.parse(xml_file).getroot()
                
                theme_info = {
                    'theme_file': theme_files[0],
                    'color_scheme': {},
                    'font_scheme': {}
                }
                
                # Extract color scheme
                color_scheme = root.find('.//a:clrScheme', self.namespaces)
                if color_scheme is not None:
                    for color_elem in color_scheme:
                        color_name = color_elem.tag.split('}')[-1] if '}' in color_elem.tag else color_elem.tag
                        theme_info['color_scheme'][color_name] = color_name
                
                # Extract font scheme
                font_scheme = root.find('.//a:fontScheme', self.namespaces)
                if font_scheme is not None:
                    major_font = font_scheme.find('.//a:majorFont/a:latin', self.namespaces)
                    minor_font = font_scheme.find('.//a:minorFont/a:latin', self.namespaces)
                    
                    if major_font is not None:
                        theme_info['font_scheme']['major_font'] = major_font.get('typeface')
                    if minor_font is not None:
                        theme_info['font_scheme']['minor_font'] = minor_font.get('typeface')
                
                return theme_info
                
        except Exception as e:
            print(f"Warning: Could not parse theme information: {e}")
            return None
    
    def _parse_numbering_xml(self, zip_ref: zipfile.ZipFile) -> Optional[Dict[str, Any]]:
        """Parse word/numbering.xml for numbering and list styles."""
        try:
            if 'word/numbering.xml' not in zip_ref.namelist():
                return None
            
            with zip_ref.open('word/numbering.xml') as xml_file:
                root = ET.parse(xml_file).getroot()
                
                numbering_info = {
                    'abstract_numbering': {},
                    'numbering_instances': {},
                    'list_styles_count': 0
                }
                
                # Parse abstract numbering definitions
                abstract_nums = root.findall('.//w:abstractNum', self.namespaces)
                numbering_info['list_styles_count'] = len(abstract_nums)
                
                for abstract_num in abstract_nums:
                    abstract_id = abstract_num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                    if abstract_id:
                        numbering_info['abstract_numbering'][abstract_id] = {
                            'abstract_num_id': abstract_id,
                            'levels': len(abstract_num.findall('.//w:lvl', self.namespaces))
                        }
                
                # Parse numbering instances
                nums = root.findall('.//w:num', self.namespaces)
                for num in nums:
                    num_id = num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                    abstract_num_ref = num.find('.//w:abstractNumId', self.namespaces)
                    if num_id and abstract_num_ref is not None:
                        numbering_info['numbering_instances'][num_id] = {
                            'num_id': num_id,
                            'abstract_num_id': abstract_num_ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        }
                
                return numbering_info
                
        except Exception as e:
            print(f"Warning: Could not parse numbering information: {e}")
            return None
    
    def save_styles(self, styles: Dict[str, Any], output_path: str) -> None:
        """
        Save extracted style information to a JSON file.
        
        Args:
            styles: The styles dictionary to save
            output_path: Path where to save the styles JSON file
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(styles, f, indent=2, ensure_ascii=False, default=str)
            print(f"Styles saved to {output_path}")
        except Exception as e:
            print(f"Error saving styles: {e}")
    
    def get_style_summary(self, styles: Dict[str, Any]) -> Dict[str, Any]:
        """
        Create a summary of the most important style information.
        
        Args:
            styles: Full styles dictionary
            
        Returns:
            Summary dictionary with key style information
        """
        summary = {}
        
        # Font summary
        if 'fonts' in styles:
            fonts = styles['fonts']
            summary['fonts'] = {
                'most_used_fonts': fonts.get('most_used_fonts', [])[:5],
                'common_sizes': fonts.get('most_common_sizes', [])[:5],
                'total_fonts_used': len(fonts.get('usage_count', {}))
            }
        
        # Color summary
        if 'colors' in styles:
            colors = styles['colors']
            summary['colors'] = {
                'color_palette_size': len(colors.get('color_palette', [])),
                'most_used_colors': colors.get('most_used_colors', [])[:5]
            }
        
        # Layout summary
        if 'layout_info' in styles and 'sections' in styles['layout_info']:
            sections = styles['layout_info']['sections']
            summary['layout'] = {
                'section_count': len(sections),
                'orientations': list(set(s.get('orientation') for s in sections if s.get('orientation'))),
                'has_headers_footers': any(s.get('different_first_page_header_footer') for s in sections)
            }
        
        # Style definitions summary
        summary['style_definitions'] = {
            'paragraph_styles_count': len(styles.get('paragraph_styles', {})),
            'character_styles_count': len(styles.get('character_styles', {})),
            'table_styles_count': len(styles.get('table_styles', {})),
            'numbering_styles_count': styles.get('numbering_styles', {}).get('list_styles_count', 0)
        }
        
        return summary